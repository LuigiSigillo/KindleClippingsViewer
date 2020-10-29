import json
from docx import Document
import wmi
import os
from shutil import copy2


def write_data(filename, songs):
    with open(filename, 'w', encoding="utf-8") as fp:
        json.dump(songs, fp, sort_keys=True, indent=4, ensure_ascii=False)


def alternative(path_to_clippings='website/files/My Clippings.txt'):
    with open(path_to_clippings, encoding='utf8') as openfileobject:
        tit, evid, empty, hlight = False, False, False, False
        diz = {}
        for line in openfileobject:
            if not tit:
                titolo = line.strip("\n")
                titolo = titolo.strip("﻿")
                
                tit = True
                try:
                    a = len(diz[titolo])
                except:
                    diz[titolo] = []
            elif not evid:
                tua_evid = line
                evid = True
            elif not empty:
                empty = True
                empt_line = line
            else:
                if not hlight:
                    hlight = True
                    highlight = ""
                if (not line.startswith("=====")):
                    highlight = highlight + line
                else:
                    tit, evid, empty, hlight = False, False, False, False
                    final_highlight = highlight.strip("\n")
                    if final_highlight == "" and len(diz[titolo]) == 0:
                            del diz[titolo]
                            continue
                    if not final_highlight in diz[titolo]: #to avoid duplicates
                        #diz[titolo].append(final_highlight)
                        is_ = is_pres_incomplete(diz[titolo], final_highlight)
                        if is_[0]: # if is present incomplete, first remove it then re-add complete
                            diz[titolo].remove(is_[1])
                        diz[titolo].append(final_highlight)
        write_data("website/files/JSONclippings.json", diz)
    return diz


def removekey(d, key):
    r = dict(d)
    del r[key]
    return r

# good if [a,ab,abc] not [a,abc,ab]. To solve this there is the second if.
def is_pres_incomplete(lista,cit):
    for h in lista:
        if cit.startswith(h) and len(cit)>len(h):
            return (True, h)
        if h[:len(cit)] == cit:
            return (True, h)
    return (False, "")
        

def write_docx(citations):
    document = Document()
    document.add_heading('Clippings', 0)
    for book in citations:
        document.add_heading(book, level=1)
        for cit in citations[book]:
            #p = document.add_paragraph(cit)
            document.add_paragraph(cit, style='List Bullet')
    
    document.add_paragraph('first item in unordered list', style='List Bullet')
    #document.add_picture('monty-truth.png', width=Inches(1.25))

    document.add_page_break()
    document.save('website/files/citations.docx')


c = wmi.WMI()

while(True):
    for disk in c.Win32_LogicalDisk():
        if(disk.volumename.startswith("Kindle")):
            letter = disk.caption
            filepath = disk.caption+"/documents/My Clippings.txt"
            copy2(filepath, "website/files/My Clippings.txt")
            citations = alternative()
            write_docx(citations)
            flag = True     
    if flag:
        break

os.system("start chrome 127.0.0.1:8000")
os.system("python -m http.server --directory website")


    

